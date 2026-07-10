import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
import io

st.set_page_config(page_title="Preschool Lesson Planner", layout="wide")
st.title("🧸 Preschool Weekly Lesson Planner")
st.markdown("**Hope Haven Preschool • Creative Curriculum + Heggerty**")

# Themes & Mighty Minutes ideas (expandable)
creative_themes = ["Exploring Balls", "Buildings", "Trees", "Gardening", "Music Making", "Reduce, Reuse, Recycle", "Pets", "Roads", "Simple Machines", "Insects", "Clothes", "Transportation", "All About Me"]
pocket_topics = ["All About Me", "Transportation", "Community Helpers", "Farm", "Ocean", "Zoo Animals", "Dinosaurs", "Space", "Seasons", "Weather", "Fall", "Winter", "Spring", "Summer", "Pirates", "Construction"]

# Mighty Minutes suggestions by theme (sample - can be expanded)
mighty_minutes = {
    "Exploring Balls": ["Mighty Minutes: Roll & Count", "Mighty Minutes: Ball Toss & Name", "Mighty Minutes: Bounce & Say", "Mighty Minutes: Sorting Balls"],
    "Buildings": ["Mighty Minutes: Block Patterns", "Mighty Minutes: Tower Challenge", "Mighty Minutes: Building Words", "Mighty Minutes: Shape Hunt"],
    "All About Me": ["Mighty Minutes: Body Parts Song", "Mighty Minutes: My Name Game", "Mighty Minutes: Feelings Charades", "Mighty Minutes: Mirror Moves"],
    # Add more as needed
}

st.sidebar.header("Week Settings")
start_date = st.sidebar.date_input("Week Start Date", datetime.now().date())
days_off_input = st.sidebar.text_area("Days Off (YYYY-MM-DD, comma separated)", placeholder="2026-07-15,2026-07-16")
theme = st.sidebar.selectbox("Main Theme (Creative Curriculum)", creative_themes)

story_for_week = st.sidebar.text_input("Storybook for the Week", placeholder="e.g., The Very Hungry Caterpillar")
pocket_topic = st.sidebar.selectbox("Pocket of Preschool Activities", pocket_topics)

if st.sidebar.button("Generate Weekly Plan"):
    off_dates = set()
    if days_off_input.strip():
        try:
            off_dates = {datetime.strptime(d.strip(), "%Y-%m-%d").date() for d in days_off_input.split(",")}
        except:
            st.error("Invalid date format")

    plan_days = []
    current = start_date
    while len(plan_days) < 4:
        if current.weekday() <= 3:
            plan_days.append(current)
        current += timedelta(days=1)

    # Mighty Minutes selection
    mm_list = mighty_minutes.get(theme, ["Mighty Minutes: Movement Activity", "Mighty Minutes: Counting Game", "Mighty Minutes: Song & Rhyme", "Mighty Minutes: Quick Transition"])
    
    rows = []
    for i, day in enumerate(plan_days):
        date_str = day.strftime("%m.%d.%Y")
        is_off = day in off_dates
        gym = "Ride tricycles" if day.weekday() == 2 else "Obstacle course / Ball skills"
        mm = mm_list[i % len(mm_list)]

        if is_off:
            rows.append({
                "Day": f"{date_str} (NO SCHOOL)",
                "Circle Time": "NO SCHOOL",
                "Gym": "NO SCHOOL",
                "Story Time": "NO SCHOOL",
                "Mighty Minutes": "NO SCHOOL",
                "Heggerty": "NO SCHOOL",
                "Math Lesson": "NO SCHOOL",
                "Language Lesson": "NO SCHOOL"
            })
        else:
            rows.append({
                "Day": date_str,
                "Circle Time": f"Morning meeting, calendar, songs - {theme}",
                "Gym": gym,
                "Story Time": story_for_week or f"Read-aloud + discussion - {theme}",
                "Mighty Minutes": mm,
                "Heggerty": "Daily 10-min phonemic awareness",
                "Math Lesson": f"Counting & number concepts - {theme}",
                "Language Lesson": f"Vocabulary & shared writing - {theme}"
            })

    df = pd.DataFrame(rows)
    st.dataframe(df, use_container_width=True, hide_index=True)

    st.subheader("Special Activities")
    st.write(f"**Weekly Art Project**: {theme} themed activity")
    st.write(f"**Pocket of Preschool**: {pocket_topic} centers & ideas")

    st.session_state.plan_days = plan_days
    st.session_state.df = df
    st.session_state.theme = theme
    st.session_state.pocket_topic = pocket_topic

# Download section (same as before, but with new columns)
if st.session_state.get("plan_days") is not None:
    if st.button("📥 Download as Word Document"):
        doc = Document()
        header_table = doc.add_table(rows=1, cols=2)
        header_table.cell(0, 0).text = "Hope Haven Preschool"
        header_table.cell(0, 1).text = "Mrs. Annette’s Class"
        
        doc.add_heading(f"Weekly Lesson Plan - {st.session_state.theme}", 0)
        doc.add_paragraph(f"Week of: {st.session_state.plan_days[0].strftime('%m.%d.%Y')} to {st.session_state.plan_days[-1].strftime('%m.%d.%Y')}")
        
        table = doc.add_table(rows=1, cols=8)
        headers = ["Day", "Circle Time", "Gym", "Story Time", "Mighty Minutes", "Heggerty", "Math Lesson", "Language Lesson"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        
        for _, row in st.session_state.df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        
        doc.add_heading("Special Activities", level=1)
        doc.add_paragraph(f"Weekly Art Project: {st.session_state.theme} themed activity")
        doc.add_paragraph(f"Pocket of Preschool: {st.session_state.pocket_topic}")
        
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.download_button(
            label="✅ Download .docx",
            data=bio.getvalue(),
            file_name=f"lesson_plan_{st.session_state.plan_days[0].strftime('%m.%d.%Y')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
